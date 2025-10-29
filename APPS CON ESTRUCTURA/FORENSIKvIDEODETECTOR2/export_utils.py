"""
Modul utilitas untuk ekspor laporan dalam berbagai format
Mengatasi masalah kompatibilitas dan memastikan ekspor berjalan lancar
"""

import os
import io
import base64
import traceback
from pathlib import Path
from datetime import datetime
from typing import Any, Dict, Optional

# Fallback jika pustaka tidak tersedia
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx tidak tersedia. Ekspor DOCX akan dinonaktifkan.")

try:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.pagesizes import letter
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("Warning: reportlab tidak tersedia. Ekspor PDF akan dinonaktifkan.")

def create_docx_report_robust(result: Any, output_path: Path) -> Optional[Path]:
    """
    Membuat laporan DOCX dengan penanganan error yang kuat
    """
    if not DOCX_AVAILABLE:
        print("Ekspor DOCX tidak tersedia karena pustaka python-docx tidak ditemukan")
        return None
    
    try:
        doc = Document()
        
        # Judul dokumen
        title = doc.add_heading('Laporan Analisis Forensik Video', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Informasi dasar
        doc.add_paragraph(f'Tanggal Analisis: {datetime.now().strftime("%d %B %Y")}')
        doc.add_paragraph(f'Nama Video: {getattr(result, "video_path", "Tidak diketahui")}')
        
        if hasattr(result, 'preservation_hash'):
            doc.add_paragraph(f'Hash Preservasi (SHA-256): {result.preservation_hash}')
        
        # Statistik utama
        doc.add_heading('Statistik Utama', level=1)
        summary = getattr(result, 'summary', {})
        doc.add_paragraph(f'Total Frame: {summary.get("total_frames", "N/A")}')
        doc.add_paragraph(f'Total Anomali: {summary.get("total_anomaly", "N/A")}')
        doc.add_paragraph(f'Persentase Anomali: {summary.get("pct_anomaly", "N/A")}%')
        
        # Informasi metadata jika tersedia
        if hasattr(result, 'metadata') and result.metadata:
            doc.add_heading('Metadata Video', level=1)
            for category, items in result.metadata.items():
                doc.add_heading(category, level=2)
                if isinstance(items, dict):
                    for key, value in items.items():
                        doc.add_paragraph(f'{key}: {value}')
                else:
                    doc.add_paragraph(str(items))
        
        # Simpan dokumen
        doc.save(str(output_path))
        return output_path
        
    except Exception as e:
        print(f"Error saat membuat laporan DOCX: {e}")
        traceback.print_exc()
        return None

def create_docx_report_in_memory(result: Any) -> Optional[bytes]:
    """
    Membuat laporan DOCX di memori dan mengembalikannya sebagai bytes.
    """
    if not DOCX_AVAILABLE:
        print("Ekspor DOCX tidak tersedia karena pustaka python-docx tidak ditemukan")
        return None
    
    try:
        doc = Document()
        
        # Judul dokumen
        title = doc.add_heading('Laporan Analisis Forensik Video', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Informasi dasar
        doc.add_paragraph(f'Tanggal Analisis: {datetime.now().strftime("%d %B %Y")}')
        doc.add_paragraph(f'Nama Video: {getattr(result, "video_path", "Tidak diketahui")}')
        
        if hasattr(result, 'preservation_hash'):
            doc.add_paragraph(f'Hash Preservasi (SHA-256): {result.preservation_hash}')
        
        # Statistik utama
        doc.add_heading('Statistik Utama', level=1)
        summary = getattr(result, 'summary', {})
        doc.add_paragraph(f'Total Frame: {summary.get("total_frames", "N/A")}')
        doc.add_paragraph(f'Total Anomali: {summary.get("total_anomaly", "N/A")}')
        doc.add_paragraph(f'Persentase Anomali: {summary.get("pct_anomaly", "N/A")}%')
        
        # Informasi metadata jika tersedia
        if hasattr(result, 'metadata') and result.metadata:
            doc.add_heading('Metadata Video', level=1)
            for category, items in result.metadata.items():
                doc.add_heading(category, level=2)
                if isinstance(items, dict):
                    for key, value in items.items():
                        doc.add_paragraph(f'{key}: {value}')
                else:
                    doc.add_paragraph(str(items))

        # Simpan dokumen ke buffer memori
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        print(f"Error saat membuat laporan DOCX di memori: {e}")
        traceback.print_exc()
        return None

def create_pdf_report_robust(result: Any, output_path: Path) -> Optional[Path]:
    """
    Membuat laporan PDF dengan penanganan error yang kuat
    """
    if not REPORTLAB_AVAILABLE:
        print("Ekspor PDF tidak tersedia karena pustaka reportlab tidak ditemukan")
        return None
    
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Judul
        title = Paragraph('Laporan Analisis Forensik Video', styles['Title'])
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Informasi dasar
        story.append(Paragraph(f'Tanggal Analisis: {datetime.now().strftime("%d %B %Y")}', styles['Normal']))
        story.append(Paragraph(f'Nama Video: {getattr(result, "video_path", "Tidak diketahui")}', styles['Normal']))
        
        if hasattr(result, 'preservation_hash'):
            story.append(Paragraph(f'Hash Preservasi (SHA-256): {result.preservation_hash}', styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Statistik utama
        story.append(Paragraph('Statistik Utama', styles['Heading1']))
        summary = getattr(result, 'summary', {})
        story.append(Paragraph(f'Total Frame: {summary.get("total_frames", "N/A")}', styles['Normal']))
        story.append(Paragraph(f'Total Anomali: {summary.get("total_anomaly", "N/A")}', styles['Normal']))
        story.append(Paragraph(f'Persentase Anomali: {summary.get("pct_anomaly", "N/A")}%', styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Bangun PDF
        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        # Simpan ke file
        with open(output_path, 'wb') as f:
            f.write(pdf_bytes)
            
        return output_path
        
    except Exception as e:
        print(f"Error saat membuat laporan PDF: {e}")
        traceback.print_exc()
        return None

def image_to_base64(image_path: str) -> Optional[str]:
    """
    Mengonversi gambar ke string base64 untuk embedding dalam laporan
    """
    try:
        if not os.path.exists(image_path):
            return None
            
        with open(image_path, "rb") as image_file:
            encoded = base64.b64encode(image_file.read()).decode('utf-8')
            
        # Tentukan tipe mime
        ext = Path(image_path).suffix.lower()
        mime_type = {
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.gif': 'image/gif'
        }.get(ext, 'image/png')
        
        return f"data:{mime_type};base64,{encoded}"
    except Exception as e:
        print(f"Error mengonversi gambar ke base64: {e}")
        return None

# Fungsi fallback jika ekspor utama gagal
def export_analysis_fallback(result: Any, output_dir: Path) -> Dict[str, Optional[Path]]:
    """
    Fallback untuk ekspor analisis dalam format sederhana
    """
    outputs = {}
    
    try:
        # Ekspor ke JSON
        json_path = output_dir / "analysis_report.json"
        # Implementasi ekspor JSON jika diperlukan
        outputs['json'] = json_path
        
        # Ekspor ringkasan ke teks
        txt_path = output_dir / "analysis_summary.txt"
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write("LAPORAN ANALISIS FORENSIK VIDEO\n")
            f.write("=" * 40 + "\n")
            f.write(f"Tanggal: {datetime.now().strftime('%d %B %Y')}\n")
            f.write(f"Video: {getattr(result, 'video_path', 'Tidak diketahui')}\n\n")
            
            summary = getattr(result, 'summary', {})
            f.write("STATISTIK UTAMA:\n")
            f.write(f"  Total Frame: {summary.get('total_frames', 'N/A')}\n")
            f.write(f"  Total Anomali: {summary.get('total_anomaly', 'N/A')}\n")
            f.write(f"  Persentase Anomali: {summary.get('pct_anomaly', 'N/A')}\n")
        outputs['txt'] = txt_path
        
    except Exception as e:
        print(f"Error dalam fallback ekspor: {e}")
        traceback.print_exc()
    
    return outputs

# Kompatibilitas backward
create_docx_report = create_docx_report_robust
create_pdf_report = create_pdf_report_robust