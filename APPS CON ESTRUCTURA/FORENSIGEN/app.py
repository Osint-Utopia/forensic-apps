import gradio as gr
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from groq import Groq  # Groq client
import unicodedata  # ✅ For cleaning Unicode text
import tempfile

# Load Groq API key from environment
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
client = Groq(api_key=GROQ_API_KEY)

# 🔧 Prompt builder
def build_prompt(images, client_name, location, date, report_type, notes):
    img_note = "Images of damage were uploaded." if images else "No images provided."
    if report_type == "Statement of Work":
        prompt = (
            f"You are a structural forensic engineer.\n\n"
            f"Generate a detailed {report_type} based on the following:\n"
            f"- Client Name: {client_name}\n"
            f"- Location: {location}\n"
            f"- Inspection Date: {date}\n"
            f"- Observations: {notes}\n"
            f"- {img_note} Observe the images and identify the nature of damage from the images.\n\n"
            f"Write in a professional tone suitable for insurance claims and contractors.\n"
            f"Organize into sections: A. Instructions, B. General Notes, C. Relevant Australian Standards, "
            f"D. Preliminaries, E. Demolition, F. Remediation and Reconstruction, G. Appendix A – Photographic Commentary."
        )
        return prompt
    elif report_type == "Causation Report":
        prompt = (
            f"You are a structural forensic engineer.\n\n"
            f"Generate a detailed {report_type} based on the following:\n"
            f"- Client Name: {client_name}\n"
            f"- Location: {location}\n"
            f"- Inspection Date: {date}\n"
            f"- Observations: {notes}\n"
            f"- {img_note} Observe the images and identify the nature of damage from the images.\n\n"
            f"Write in a professional tone suitable for insurance claims and contractors.\n"
            f"Organize into sections: A. Instructions (State the commissioning entity, the purpose of the inspection, and the date of site attendance), "
            f"B. General Notes (Briefly describe the nature of the damage), C. Observations (Provide a descriptive assessment of the damage visible in the supplied photographs and the effect of this damage on the structural integritry of the structure), "
            f"D. Discussion, E. Conclusion, F. Limitations to the Report, G. Appendix A – Photographic Commentary."
        )
        return prompt
    else:
        prompt = (
            f"You are a structural forensic engineer.\n\n"
            f"Generate a detailed {report_type} based on the following:\n"
            f"- Client Name: {client_name}\n"
            f"- Location: {location}\n"
            f"- Inspection Date: {date}\n"
            f"- Observations: {notes}\n"
            f"- {img_note} Observe the images and identify the nature of damage from the images.\n\n"
            f"Write in a professional tone suitable for insurance claims and contractors.\n"
            f"Organize into sections: A. Instructions (State the commissioning entity, the purpose of the inspection, and the date of site attendance), "
            f"B. General Notes (Briefly describe the nature of the damage), C. Observations (Provide a descriptive assessment of the damage visible in the supplied photographs and the effect of this damage on the structural integritry of the structure), "
            f"D. Discussion, E. Conclusion, F. Limitations to the Report, G. Appendix A – Photographic Commentary."
        )
        return prompt
# 🤖 Groq-based report generation
def generate_report(images, client_name, location, date, report_type, notes):
    if not all([client_name, location, date, report_type]):
        return "⚠️ Please fill out all fields."
    
    prompt = build_prompt(images, client_name, location, date, report_type, notes)
    
    try:
        response = client.chat.completions.create(
            model="llama3-70b-8192",
            messages=[
                {"role": "system", "content": "You are a structural forensic engineering assistant."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.4,
            max_tokens=1024
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"❌ Error generating report: {str(e)}"

# 🧼 Clean Unicode text
def clean_text(text):
    # Normalize and strip non-ASCII characters
    return unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")
    
def add_table_of_contents(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')  # Create begin field
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')  # Field code
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # TOC field code

    fldChar2 = OxmlElement('w:fldChar')  # Separate field
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')  # End field
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    
# Export to DOCX with sections, images, and header/footer
def export_to_docx(text, images, client_name, date, logo_path=None, add_watermark=False):
    text = clean_text(text)
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    section = doc.sections[0]

    # 📌 HEADER: Logo + Metadata
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.0))  # or 6.5 based on your layout
    table.allow_autofit = True
    row = table.rows[0]
    cell_logo, cell_info = row.cells

    # 🖼 Insert company logo if available
    if logo_path and os.path.exists(logo_path):
        try:
            paragraph = cell_logo.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(logo_path, width=Inches(1.0))
        except Exception:
            cell_logo.text = "[Logo]"

    # 🧾 Add client info to header
    info_paragraph = cell_info.paragraphs[0]
    info_paragraph.text = f"Forensic Report\nClient: {client_name}\nInspection Date: {date}"
    info_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # === TOC ===
    doc.add_paragraph("Table of Contents", style='Heading 1')
    toc_paragraph = doc.add_paragraph()
    add_table_of_contents(toc_paragraph)
    doc.add_page_break()

    # 🏷 TITLE
    doc.add_paragraph().add_run().add_break()  # spacer
    title = doc.add_heading("RMA Forensic Engineering \n{report_type}", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 📚 Section Parsing (A. to G.)
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line[:2] in ['A.', 'B.', 'C.', 'D.', 'E.', 'F.', 'G.']:
            doc.add_heading(line, level=2)
        else:
            para = doc.add_paragraph(line)
            para.paragraph_format.space_after = Pt(6)

    # 📷 Appendix – Images
    if images:
        doc.add_page_break()
        doc.add_heading("Appendix A – Photographic Commentary", level=2)
        for i, img in enumerate(images, 1):
            try:
                img_path = img.name if hasattr(img, "name") else img
                doc.add_picture(img_path, width=Inches(5.5))
                caption = doc.add_paragraph(f"Figure {i}: {os.path.basename(img_path).replace('_', ' ').split('.')[0].title()}")
                caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                caption.runs[0].bold = True
                doc.add_paragraph()
            except Exception as e:
                doc.add_paragraph(f"[Error inserting image: {str(e)}]")

    # ➕ Footer
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = f"Generated by ForensiGen – {client_name} | {date}"
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_paragraph.runs[0].italic = True

    # === PAGE NUMBERS ===
    footer_run = footer_paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    footer_run._r.append(fldChar)
    footer_run._r.append(instrText)
    footer_run._r.append(fldChar2)
    footer_run._r.append(fldChar3)

    # === OPTIONAL: Watermark (Text-only) ===
    if add_watermark:
        watermark = section.header.paragraphs[0].add_run("CONFIDENTIAL")
        watermark.font.size = Pt(40)
        watermark.font.color.rgb = RGBColor(200, 200, 200)

    # 📂 Save file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name


# 🎛️ Gradio UI
with gr.Blocks() as demo:
    gr.Markdown("# 🏚️ ForensiGen – AI Forensic Report Generator")

    with gr.Row():
        images = gr.File(label="Upload Damage Images", file_types=["image"], file_count="multiple")
        report_type = gr.Dropdown(choices=["Statement of Work", "Causation Report", "Structural Integrity Report"], label="Report Type")

    with gr.Row():
        client_name = gr.Textbox(label="Client Name")
        location = gr.Textbox(label="Location")
        date = gr.Textbox(label="Inspection Date (e.g., 2025-05-17)")

    notes = gr.Textbox(label="Additional Observations", lines=4, placeholder="Describe what was observed on-site... (Optional)")

    logo_path = gr.Textbox(value="assets/logo.png", visible=False)
    watermark_toggle = gr.Checkbox(label="Add Confidential Watermark")

    generate_btn = gr.Button("🔍 Generate Report")
    report_output = gr.Textbox(label="Generated Report", lines=20)
    download_btn = gr.Button("⬇️ Export as DOCX")
    file_output = gr.File(label="Download DOCX", file_types=[".docx"])

    generate_btn.click(fn=generate_report,
                       inputs=[images, client_name, location, date, report_type, notes],
                       outputs=report_output)

    download_btn.click(fn=export_to_docx,
                   inputs=[report_output, images, client_name, date, logo_path, watermark_toggle],
                   outputs=file_output)

demo.launch()
