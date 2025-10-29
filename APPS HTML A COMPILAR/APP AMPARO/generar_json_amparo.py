# -*- coding: utf-8 -*-
"""
Este script combina los artículos de la Ley de Amparo y las tesis jurisprudenciales
para generar un archivo JSON con la estructura especificada.
Requiere: python-docx instalado (pip install python-docx)
"""

import json
import re
from docx import Document

# Cargar documentos
doc_ley = Document("ley de amparo para convinar.docx")
doc_tesis = Document("tesis de amparo para convinar-.docx")

# Función para extraer texto plano de un DOCX
def extraer_texto(doc):
    return "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

texto_ley = extraer_texto(doc_ley)
texto_tesis = extraer_texto(doc_tesis)

# Extraer artículos
def extraer_articulos(texto):
    articulos = {}
    titulos = []
    capitulos = []
    actual = {}
    numero = None

    for linea in texto.splitlines():
        linea = linea.strip()
        if not linea:
            continue
        if linea.upper().startswith("TÍTULO "):
            titulos.append(linea)
        elif linea.upper().startswith("CAPÍTULO "):
            capitulos.append(linea)
        elif re.match(r"^Artículo\s+\d+[oº.]?", linea, re.IGNORECASE):
            if numero:
                articulos[numero] = actual
            numero = int(re.search(r"\d+", linea).group())
            actual = {
                "titulo": titulos[-1] if titulos else "",
                "capitulo": capitulos[-1] if capitulos else "",
                "numero": numero,
                "texto": "",
                "fracciones": [],
                "texto_adicional": "",
                "articulos_relacionados": [],
                "tema_general": "",
                "tesis_relacionadas": []
            }
        elif re.match(r"^[IVXLCDM]+\.\s", linea):
            actual["fracciones"].append(linea)
        elif numero:
            if not actual["texto"]:
                actual["texto"] = linea
            else:
                actual["texto_adicional"] += " " + linea

    if numero:
        articulos[numero] = actual

    return articulos

# Extraer tesis
def extraer_tesis(texto):
    bloques = re.split(r"ART[IÍ]CULO\s+(\d+)[.]*", texto)
    tesis_por_art = {}
    for i in range(1, len(bloques), 2):
        num = int(bloques[i])
        contenido = bloques[i+1].strip()
        tesis = re.split(r"\n(?=[A-ZÁÉÍÓÚÑ ]+?\.)", contenido)
        tesis = [t.strip().replace('\n', ' ') for t in tesis if t.strip()]
        tesis_por_art[num] = tesis
    return tesis_por_art

# Procesar datos
articulos = extraer_articulos(texto_ley)
tesis_raw = extraer_tesis(texto_tesis)
tesis_json = {}

for art_num, lista in tesis_raw.items():
    for idx, contenido in enumerate(lista, 1):
        clave = f"art{art_num}_tesis{idx}"
        tesis_json[clave] = {
            "articulo": art_num,
            "rubro": contenido.split('.')[0][:180],
            "registro": "s/nr",
            "texto": contenido,
            "materia": "",
            "organo": "",
            "epoca": ""
        }
        if art_num in articulos:
            articulos[art_num]["tesis_relacionadas"].append(clave)

# Construir JSON final
resultado = {
    "articulos": articulos,
    "tesis": tesis_json,
    "configuracion": {
        "max_tesis_por_articulo": max(len(v["tesis_relacionadas"]) for v in articulos.values()),
        "total_articulos": len(articulos),
        "version": "1.0"
    }
}

# Guardar archivo
with open("ley_amparo_tesis_completo.json", "w", encoding="utf-8") as f:
    json.dump(resultado, f, ensure_ascii=False, indent=2)

print("✅ Archivo JSON generado: ley_amparo_tesis_completo.json")
